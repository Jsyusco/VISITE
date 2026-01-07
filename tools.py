# utils.py
import streamlit as st
import pandas as pd
import numpy as np
import uuid
import json
import zipfile
import io
import urllib.parse
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
[cite_start]from streamlit_gsheets import GSheetsConnection # [cite: 1]

# --- CONSTANTES ---
PROJECT_RENAME_MAP = {
    'Intitulé': 'Intitulé',
    'Fournisseur Bornes AC [Bornes]': 'Fournisseur Bornes AC',
    'Fournisseur Bornes DC [Bornes]': 'Fournisseur Bornes DC',
    'L [Plan de Déploiement]': 'PDC Lent',
    'R [Plan de Déploiement]': 'PDC Rapide',
    'UR [Plan de Déploiement]': 'PDC Ultra-rapide',
    'Pré L [Plan de Déploiement]': 'PDC L pré-équipés',
    'Pré R [Plan de Déploiement]': 'PDC R pré-équipés',
    'Pré UR [Plan de Déploiement]': 'PDC UR pré-équipés',
}

DISPLAY_GROUPS = [
    ['Intitulé', 'Fournisseur Bornes AC [Bornes]', 'Fournisseur Bornes DC [Bornes]'],
    ['L [Plan de Déploiement]', 'R [Plan de Déploiement]', 'UR [Plan de Déploiement]'],
    ['Pré L [Plan de Déploiement]', 'Pré R [Plan de Déploiement]', 'Pré UR [Plan de Déploiement]'],
]

SECTION_PHOTO_RULES = {
    "Bornes DC": ['R [Plan de Déploiement]', 'UR [Plan de Déploiement]'],
    "Bornes AC": ['L [Plan de Déploiement]'],
}

COMMENT_ID = 100
COMMENT_QUESTION = "Veuillez préciser pourquoi le nombre de photo partagé ne correspond pas au minimum attendu"

# --- CONNEXION GOOGLE SHEETS ---
def get_db_connection():
    return st.connection("gsheets", type=GSheetsConnection)

# --- CHARGEMENT DONNÉES ---
@st.cache_data(ttl=600)
def load_form_structure_from_sheets():
    try:
        conn = get_db_connection()
        # Lecture de l'onglet 'Questions'
        df = conn.read(worksheet="Questions") 
        
        # Nettoyage des colonnes
        df.columns = df.columns.str.strip()
        
        rename_map = {'Conditon value': 'Condition value', 'condition value': 'Condition value', 'Condition Value': 'Condition value', 'Condition': 'Condition value', 'Conditon on': 'Condition on', 'condition on': 'Condition on'}
        actual_rename = {k: v for k, v in rename_map.items() if k in df.columns}
        df = df.rename(columns=actual_rename)
        
        expected_cols = ['options', 'Description', 'Condition value', 'Condition on', 'section', 'id', 'question', 'type', 'obligatoire']
        for col in expected_cols:
            if col not in df.columns: df[col] = np.nan 
        
        df['options'] = df['options'].fillna('')
        df['Description'] = df['Description'].fillna('')
        df['Condition value'] = df['Condition value'].fillna('')
        df['Condition on'] = pd.to_numeric(df['Condition on'], errors='coerce').fillna(0).astype(int)
        
        for col in df.select_dtypes(include=['object']).columns:
            df[col] = df[col].astype(str).str.strip()
        return df
    except Exception as e:
        st.error(f"Erreur chargement structure (Sheet 'Questions'): {e}")
        return None

@st.cache_data(ttl=3600)
def load_site_data_from_sheets():
    try:
        conn = get_db_connection()
        # Lecture de l'onglet 'Sites'
        df_site = conn.read(worksheet="Sites")
        df_site.columns = df_site.columns.str.strip()
        return df_site
    except Exception as e:
        st.error(f"Erreur chargement sites (Sheet 'Sites'): {e}")
        return None

# --- LOGIQUE MÉTIER (Inchangée) ---
def get_expected_photo_count(section_name, project_data):
    if section_name.strip() not in SECTION_PHOTO_RULES:
        return None, None 

    columns = SECTION_PHOTO_RULES[section_name.strip()]
    total_expected = 0
    details = []

    for col in columns:
        val = project_data.get(col, 0)
        try:
            if pd.isna(val) or val == "":
                num = 0
            else:
                num = int(float(str(val).replace(',', '.'))) 
        except Exception:
            num = 0
        total_expected += num
        short_name = PROJECT_RENAME_MAP.get(col, col) 
        details.append(f"{num} {short_name}")

    detail_str = " + ".join(details)
    return total_expected, detail_str

def evaluate_single_condition(condition_str, all_answers):
    if "=" not in condition_str:
        return True
    try:
        target_id_str, expected_value_raw = condition_str.split('=', 1)
        target_id = int(target_id_str.strip())
        expected_value = expected_value_raw.strip().strip('"').strip("'")
        user_answer = all_answers.get(target_id)
        if user_answer is not None:
            return str(user_answer).strip().lower() == str(expected_value).strip().lower()
        else:
            return False
    except Exception:
        return True

def check_condition(row, current_answers, collected_data):
    try:
        if int(row.get('Condition on', 0)) != 1: return True
    except (ValueError, TypeError): return True

    all_past_answers = {}
    for phase_data in collected_data: 
        all_past_answers.update(phase_data['answers'])
    combined_answers = {**all_past_answers, **current_answers}
    
    condition_raw = str(row.get('Condition value', '')).strip().strip('"').strip("'")
    if not condition_raw: return True

    or_blocks = condition_raw.split(' OU ')
    for block in or_blocks:
        and_conditions = block.split(' ET ')
        block_is_valid = True
        for atom in and_conditions:
            if not evaluate_single_condition(atom, combined_answers):
                block_is_valid = False
                break
        if block_is_valid:
            return True
    return False

def validate_section(df_questions, section_name, answers, collected_data, project_data):
    missing = []
    section_rows = df_questions[df_questions['section'] == section_name]
    comment_val = answers.get(COMMENT_ID)
    has_justification = comment_val is not None and str(comment_val).strip() != ""
    
    expected_total_base, detail_str = get_expected_photo_count(section_name.strip(), project_data)
    expected_total = expected_total_base
    
    photo_question_count = sum(
        1 for _, row in section_rows.iterrows()
        if str(row.get('type', '')).strip().lower() == 'photo' and check_condition(row, answers, collected_data)
    )
    
    if expected_total is not None and expected_total > 0:
        expected_total = expected_total_base * photo_question_count
        detail_str = f"{detail_str} | Questions photo visibles: {photo_question_count} -> Total ajusté: {expected_total}"

    current_photo_count = 0
    photo_questions_found = False
    
    for _, row in section_rows.iterrows():
        q_type = str(row['type']).strip().lower()
        if q_type == 'photo' and check_condition(row, answers, collected_data):
            photo_questions_found = True
            q_id = int(row['id'])
            val = answers.get(q_id)
            if isinstance(val, list):
                current_photo_count += len(val)

    for _, row in section_rows.iterrows():
        q_id = int(row['id'])
        if q_id == COMMENT_ID: continue
        if not check_condition(row, answers, collected_data): continue
        is_mandatory = str(row['obligatoire']).strip().lower() == 'oui'
        q_type = str(row['type']).strip().lower()
        val = answers.get(q_id)
        if is_mandatory:
            if q_type == 'photo':
                if not isinstance(val, list) or len(val) == 0:
                    missing.append(f"Question {q_id} : {row['question']} (Au moins une photo est requise)")
            else:
                if isinstance(val, list):
                    if not val: missing.append(f"Question {q_id} : {row['question']} (fichier(s) manquant(s))")
                elif val is None or val == "" or (isinstance(val, (int, float)) and val == 0):
                    missing.append(f"Question {q_id} : {row['question']}")

    is_photo_count_incorrect = False
    if expected_total is not None and expected_total > 0:
        if photo_questions_found and current_photo_count != expected_total:
            is_photo_count_incorrect = True
            error_message = (
                f"⚠️ **Écart de Photos pour '{str(section_name)}'**.\n"
                f"Attendu : **{str(expected_total)}** (calculé : {str(detail_str)}).\n"
                f"Reçu : **{str(current_photo_count)}**.\n"
            )
            if not has_justification:
                missing.append(
                    f"**Commentaire (ID {COMMENT_ID}) :** {COMMENT_QUESTION} "
                    f"(requis en raison de l'écart de photo). \n\n {error_message}"
                )

    if not is_photo_count_incorrect and COMMENT_ID in answers:
        del answers[COMMENT_ID]

    return len(missing) == 0, missing

# --- SAUVEGARDE ET EXPORTS (Modifié pour Sheets) ---

def define_custom_styles(doc):
    try: title_style = doc.styles.add_style('Report Title', WD_STYLE_TYPE.PARAGRAPH)
    except: title_style = doc.styles['Report Title']
    title_font = title_style.font
    title_font.name, title_font.size, title_font.bold = 'Arial', Pt(20), True
    title_font.color.rgb = RGBColor(0x01, 0x38, 0x2D)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(20)

    try: subtitle_style = doc.styles.add_style('Report Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    except: subtitle_style = doc.styles['Report Subtitle']
    subtitle_font = subtitle_style.font
    subtitle_font.name, subtitle_font.size, subtitle_font.bold = 'Arial', Pt(14), True
    subtitle_font.color.rgb = RGBColor(0x00, 0x56, 0x47)
    subtitle_style.paragraph_format.space_after = Pt(10)

    try: text_style = doc.styles.add_style('Report Text', WD_STYLE_TYPE.PARAGRAPH)
    except: text_style = doc.styles['Report Text']
    text_font = text_style.font
    text_font.name, text_font.size = 'Calibri', Pt(11)
    text_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def create_word_report(collected_data, df_struct, project_data, form_start_time):
    doc = Document()
    define_custom_styles(doc)
    
    doc.add_paragraph('Rapport d\'Audit Chantier', style='Report Title')

    doc.add_paragraph('Informations du Projet', style='Report Subtitle')
    project_table = doc.add_table(rows=3, cols=2)
    project_table.style = 'Light Grid Accent 1'
    
    project_table.rows[0].cells[0].text = 'Intitulé'
    project_table.rows[0].cells[1].text = str(project_data.get('Intitulé', 'N/A'))
    
    start_time_str = form_start_time.strftime('%d/%m/%Y %H:%M') if form_start_time else "N/A"
    project_table.rows[1].cells[0].text = 'Date de début'
    project_table.rows[1].cells[1].text = start_time_str
    project_table.rows[2].cells[0].text = 'Date de fin'
    project_table.rows[2].cells[1].text = datetime.now().strftime('%d/%m/%Y %H:%M')

    for row in project_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs: p.style = 'Report Text'
    
    doc.add_paragraph()
    doc.add_paragraph('Détails du Projet', style='Report Subtitle')
    for group in DISPLAY_GROUPS:
        for field_key in group:
            renamed_key = PROJECT_RENAME_MAP.get(field_key, field_key)
            value = project_data.get(field_key, 'N/A')
            p = doc.add_paragraph(style='Report Text')
            p.add_run(f'{renamed_key}: ').bold = True
            p.add_run(str(value))
    
    doc.add_page_break()
    
    for phase_idx, phase in enumerate(collected_data):
        doc.add_paragraph(f'Phase: {phase["phase_name"]}', style='Report Subtitle')
        
        for q_id, answer in phase['answers'].items():
            if int(q_id) == COMMENT_ID:
                q_text = COMMENT_QUESTION
            else:
                q_row = df_struct[df_struct['id'].astype(int) == int(q_id)]
                q_text = q_row.iloc[0]['question'] if not q_row.empty else f"ID {q_id}"
            
            is_photo = (isinstance(answer, list) and answer and hasattr(answer[0], 'read')) or hasattr(answer, 'read')
            
            if is_photo:
                doc.add_paragraph(f'Q{q_id}: {q_text}', style='Report Subtitle')
                photos = answer if isinstance(answer, list) else [answer]
                for idx, f_obj in enumerate(photos):
                    try:
                        f_obj.seek(0)
                        doc.add_picture(BytesIO(f_obj.read()), width=Inches(5))
                        cap = doc.add_paragraph(f'Photo {idx+1}: {f_obj.name}', style='Report Text')
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if cap.runs: 
                            cap.runs[0].font.size, cap.runs[0].font.italic = Pt(9), True
                        f_obj.seek(0)
                    except: doc.add_paragraph(f"[Erreur Photo {idx+1}]", style='Report Text')
                doc.add_paragraph()
            else:
                t = doc.add_table(rows=1, cols=2)
                t.style = 'Light Grid Accent 1'
                t.cell(0,0).text = f'Q{q_id}: {q_text}'
                t.cell(0,1).text = str(answer)
                for cell in t.rows[0].cells:
                    cell.paragraphs[0].style = 'Report Text'
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                t.cell(0,0).paragraphs[0].runs[0].bold = True
                doc.add_paragraph()
        
        if phase_idx < len(collected_data) - 1: doc.add_page_break()
    
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def save_form_data(collected_data, project_data, submission_id, start_time):
    """
    Sauvegarde une nouvelle ligne dans l'onglet 'Reponses' du Google Sheet.
    Les données complexes sont sérialisées en JSON.
    NOTE: Les photos ne sont PAS uploadées dans le Sheet. Seuls les noms de fichiers sont conservés.
    """
    try:
        conn = get_db_connection()
        
        # 1. Nettoyage et conversion des données pour le JSON
        cleaned_data = []
        for phase in collected_data:
            clean_phase = {"phase_name": phase["phase_name"], "answers": {}}
            for k, v in phase["answers"].items():
                # Transformation des objets fichiers en chaînes de caractères (noms)
                if isinstance(v, list) and v and hasattr(v[0], 'read'): 
                    file_names = ", ".join([f.name for f in v])
                    clean_phase["answers"][str(k)] = f"Fichiers: {file_names}"
                elif hasattr(v, 'read'): 
                     clean_phase["answers"][str(k)] = f"Fichier: {v.name}"
                else:
                    clean_phase["answers"][str(k)] = v
            cleaned_data.append(clean_phase)
        
        json_dump = json.dumps(cleaned_data, ensure_ascii=False)
        
        # 2. Création de la ligne à insérer
        new_row = pd.DataFrame([{
            "ID": submission_id,
            "Date": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            "Projet": project_data.get('Intitulé', 'N/A'),
            "Donnees_JSON": json_dump
        }])
        
        # 3. Ajout à la feuille existante
        # Note : st-gsheets lit tout, ajoute, et réécrit.
        existing_data = conn.read(worksheet="Reponses")
        updated_df = pd.concat([existing_data, new_row], ignore_index=True)
        conn.update(worksheet="Reponses", data=updated_df)
        
        return True, submission_id 
    except Exception as e:
        return False, str(e)

def create_csv_export(collected_data, df_struct, project_name, submission_id, start_time):
    data_for_df = []
    for phase in collected_data:
        for q_id, answer in phase['answers'].items():
            if not hasattr(answer, 'read') and not (isinstance(answer, list) and answer and hasattr(answer[0], 'read')):
                data_for_df.append({
                    'Projet': project_name, 'Phase': phase['phase_name'],
                    'Question_ID': q_id, 'Réponse': answer
                })
    return pd.DataFrame(data_for_df).to_csv(index=False).encode('utf-8')

def create_zip_export(collected_data):
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w') as zip_file:
        for phase in collected_data:
            for q_id, files in phase['answers'].items():
                photos = files if isinstance(files, list) else [files]
                for i, f in enumerate(photos):
                    if hasattr(f, 'getvalue'):
                        zip_file.writestr(f"{phase['phase_name']}_Q{q_id}_{i}.jpg", f.getvalue())
    buf.seek(0)
    return buf

# --- COMPOSANT UI (Inchangé) ---
def render_question(row, answers, phase_name, key_suffix, loop_index, project_data):
    q_id = int(row.get('id', 0))
    is_dynamic_comment = (q_id == COMMENT_ID)
    
    if is_dynamic_comment:
        q_text, q_type, q_desc, q_mandatory = COMMENT_QUESTION, 'text', "Requis si écart photo.", True
        q_options = []
    else:
        q_text = row['question']
        q_type, q_desc = str(row['type']).strip().lower(), row['Description']
        q_mandatory = str(row['obligatoire']).lower() == 'oui'
        q_options = str(row['options']).split(',') if row['options'] else []

    label_html = f"<strong>{q_id}. {q_text}</strong>" + (' <span class="mandatory">*</span>' if q_mandatory else "")
    widget_key = f"q_{q_id}_{phase_name}_{key_suffix}_{loop_index}"
    current_val = answers.get(q_id)

    st.markdown(f'<div class="question-card"><div>{label_html}</div>', unsafe_allow_html=True)
    if q_desc: st.markdown(f'<div class="description">⚠️ {q_desc}</div>', unsafe_allow_html=True)
    
    if q_type == 'text':
        answers[q_id] = st.text_area("R", value=current_val if current_val else "", key=widget_key, label_visibility="collapsed") if is_dynamic_comment else st.text_input("R", value=current_val if current_val else "", key=widget_key, label_visibility="collapsed")
    elif q_type == 'select':
        opts = [o.strip() for o in q_options]
        if "" not in opts: opts.insert(0, "")
        answers[q_id] = st.selectbox("S", opts, index=opts.index(current_val) if current_val in opts else 0, key=widget_key, label_visibility="collapsed")
    elif q_type == 'number':
        answers[q_id] = st.number_input("N", value=int(current_val) if current_val else 0, step=1, key=widget_key, label_visibility="collapsed")
    elif q_type == 'photo':
        exp, det = get_expected_photo_count(phase_name.strip(), project_data)
        if exp: st.info(f"📸 **Attendu : {exp}** ({det})")
        answers[q_id] = st.file_uploader("I", type=['png', 'jpg', 'jpeg'], accept_multiple_files=True, key=widget_key, label_visibility="collapsed")
    st.markdown('</div>', unsafe_allow_html=True)
